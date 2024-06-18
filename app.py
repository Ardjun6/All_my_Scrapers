from flask import Flask, render_template_string, send_file, request
from selenium import webdriver
from selenium.webdriver.chrome.service import Service
from selenium.webdriver.common.by import By
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from bs4 import BeautifulSoup
import pandas as pd
import time
from io import BytesIO
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from webdriver_manager.chrome import ChromeDriverManager

app = Flask(__name__)

def init_webdriver():
    options = Options()
    options.add_argument('--headless')
    options.add_argument('--no-sandbox')
    options.add_argument('--disable-dev-shm-usage')

    service = Service(ChromeDriverManager().install())
    return webdriver.Chrome(service=service, options=options)

def scrape_steam_market(page):
    driver = init_webdriver()
    
    url = 'https://steamcommunity.com/market/search?appid=730' if page == 1 else f'https://steamcommunity.com/market/search?appid=730#p{page}_popular_desc'
    driver.get(url)

    try:
        WebDriverWait(driver, 10).until(EC.presence_of_element_located((By.ID, 'searchResultsRows')))
        soup = BeautifulSoup(driver.page_source, 'html.parser')
    except Exception as e:
        driver.quit()
        return []
    
    driver.quit()

    results = soup.find('div', id='searchResultsRows')
    top_items = []
    for item in results.find_all('a', class_='market_listing_row_link'):
        name = item.find('span', class_='market_listing_item_name').text
        price = item.find('span', class_='normal_price').text
        link = item['href']
        image = item.find('img')['src']
        top_items.append({
            'name': name,
            'price': price,
            'link': link,
            'image': image
        })
    return top_items

@app.route('/')
def index():
    page = int(request.args.get('page', 1))
    top_items = scrape_steam_market(page)
    
    html = '''
    <!doctype html>
    <html lang="en">
      <head>
        <meta charset="utf-8">
        <meta name="viewport" content="width=device-width, initial-scale=1, shrink-to-fit=no">
        <title>Steam Market Top Items</title>
        <link rel="stylesheet" href="https://stackpath.bootstrapcdn.com/bootstrap/4.5.2/css/bootstrap.min.css">
      </head>
      <body>
        <div class="container mt-5">
          <h1 class="mb-4">Steam Market Top Items</h1>
          <a href="/download/excel?page={{ page }}" class="btn btn-success mb-4">Download Excel</a>
          <a href="/download/csv?page={{ page }}" class="btn btn-info mb-4">Download CSV</a>
          <a href="/download/word?page={{ page }}" class="btn btn-primary mb-4">Download Word</a>
          <div class="row">
            {% for item in items %}
              <div class="col-md-6 col-lg-4 mb-4">
                <div class="card h-100">
                  <img src="{{ item.image }}" class="card-img-top" alt="{{ item.name }}">
                  <div class="card-body">
                    <h5 class="card-title">{{ item.name }}</h5>
                    <p class="card-text">Price: {{ item.price }}</p>
                    <a href="{{ item.link }}" class="btn btn-primary" target="_blank">View on Steam</a>
                  </div>
                </div>
              </div>
            {% endfor %}
          </div>
          <nav aria-label="Page navigation">
            <ul class="pagination">
              {% if page > 1 %}
              <li class="page-item">
                <a class="page-link" href="/?page={{ page-1 }}" aria-label="Previous">
                  <span aria-hidden="true">&laquo;</span>
                </a>
              </li>
              {% endif %}
              <li class="page-item"><a class="page-link" href="#">{{ page }}</a></li>
              <li class="page-item">
                <a class="page-link" href="/?page={{ page+1 }}" aria-label="Next">
                  <span aria-hidden="true">&raquo;</span>
                </a>
              </li>
            </ul>
          </nav>
        </div>
      </body>
    </html>
    '''
    
    return render_template_string(html, items=top_items, page=page)

@app.route('/download/<file_type>')
def download(file_type):
    page = int(request.args.get('page', 1))
    top_items = scrape_steam_market(page)
    
    if file_type == 'excel':
        return download_excel(top_items)
    elif file_type == 'csv':
        return download_csv(top_items)
    elif file_type == 'word':
        return download_word(top_items)
    else:
        return "Invalid file type requested.", 400

def download_excel(data):
    df = pd.DataFrame(data)
    df.columns = ['Name', 'Price', 'Link', 'Image URL']
    output = BytesIO()
    with pd.ExcelWriter(output, engine='xlsxwriter') as writer:
        df.to_excel(writer, index=False, sheet_name='Top Items')
        workbook = writer.book
        worksheet = writer.sheets['Top Items']
        worksheet.set_column('A:A', 30)  # Name column
        worksheet.set_column('B:B', 20)  # Price column
        worksheet.set_column('C:C', 50)  # Link column
        worksheet.set_column('D:D', 50)  # Image URL column

        header_format = workbook.add_format({
            'bold': True,
            'text_wrap': True,
            'valign': 'top',
            'fg_color': '#D7E4BC',
            'border': 1,
            'align': 'center'
        })
        
        for col_num, value in enumerate(df.columns.values):
            worksheet.write(0, col_num, value, header_format)
        
        row_formats = [
            workbook.add_format({'bg_color': '#FFFFFF'}),
            workbook.add_format({'bg_color': '#F3F3F3'})
        ]
        
        for row_num, row_data in enumerate(df.values, 1):
            row_format = row_formats[row_num % 2]
            for col_num, cell_data in enumerate(row_data):
                worksheet.write(row_num, col_num, cell_data, row_format)

        price_format = workbook.add_format({
            'bg_color': '#FFEB9C',
            'font_color': '#9C5700'
        })
        worksheet.conditional_format('B2:B{}'.format(len(df) + 1), {'type': 'text', 'criteria': 'containing', 'value': 'USD', 'format': price_format})

    output.seek(0)
    return send_file(output, download_name='steam_market_top_items.xlsx', as_attachment=True)

def download_csv(data):
    df = pd.DataFrame(data)
    df.columns = ['Name', 'Price', 'Link', 'Image URL']
    output = BytesIO()
    df.to_csv(output, index=False)
    output.seek(0)
    return send_file(output, download_name='steam_market_top_items.csv', as_attachment=True)

def download_word(data):
    document = Document()
    document.add_heading('Steam Market Top Items', 0).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    table = document.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Name'
    hdr_cells[1].text = 'Price'
    hdr_cells[2].text = 'Link'
    hdr_cells[3].text = 'Image URL'
    
    for item in data:
        row_cells = table.add_row().cells
        row_cells[0].text = item['name']
        row_cells[1].text = item['price']
        row_cells[2].text = item['link']
        row_cells[3].text = item['image']
    
    output = BytesIO()
    document.save(output)
    output.seek(0)
    return send_file(output, download_name='steam_market_top_items.docx', as_attachment=True)

if __name__ == '__main__':
    app.run(debug=True)
